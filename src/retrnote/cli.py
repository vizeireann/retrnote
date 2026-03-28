import sys
import argparse
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.opc import constants
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup, NavigableString
import subprocess

from . import __version__

def add_run(container, text, styles, preserve_newlines):
    if preserve_newlines:
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                br_run = OxmlElement('w:r')
                br_run.append(OxmlElement('w:br'))
                container.append(br_run)
            if line:
                run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                if styles.get('bold', False):
                    rPr.append(OxmlElement('w:b'))
                if styles.get('italic', False):
                    rPr.append(OxmlElement('w:i'))
                if styles.get('underline', False):
                    u = OxmlElement('w:u')
                    u.set(qn('w:val'), 'single')
                    rPr.append(u)
                if styles.get('hyperlink', False):
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')
                    rPr.append(color)
                    if not styles.get('underline', False):
                        u = OxmlElement('w:u')
                        u.set(qn('w:val'), 'single')
                        rPr.append(u)
                run.append(rPr)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = line
                run.append(t)
                container.append(run)
    else:
        text = re.sub(r'\s+', ' ', text)
        if not text.strip():
            return
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if styles.get('bold', False):
            rPr.append(OxmlElement('w:b'))
        if styles.get('italic', False):
            rPr.append(OxmlElement('w:i'))
        if styles.get('underline', False):
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        if styles.get('hyperlink', False):
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            if not styles.get('underline', False):
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
        run.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run.append(t)
        container.append(run)

def process_children(nodes, doc, parent_container=None, styles=None, preserve_newlines=True, justify=True):
    if styles is None:
        styles = {'bold': False, 'italic': False, 'underline': False, 'hyperlink': False}
    current_p = None
    for node in nodes:
        if isinstance(node, NavigableString):
            text = str(node)
            if not preserve_newlines and not text.strip():
                continue
            if parent_container is None:
                if current_p is None:
                    current_p = doc.add_paragraph()
                    if justify:
                        current_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    parent_container = current_p._element
            add_run(parent_container, text, styles, preserve_newlines)
        elif node.name:
            new_styles = styles.copy()
            is_block = False
            if node.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                is_block = True
                if node.name == 'p':
                    current_p = doc.add_paragraph()
                    if justify:
                        current_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    level = int(node.name[1])
                    current_p = doc.add_paragraph()
                    current_p.style = f'Heading {level}'
                    if node.has_attr('align') and node['align'].lower() == 'center':
                        current_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                parent_container = current_p._element
                process_children(node.children, doc, parent_container, new_styles, preserve_newlines, justify)
                current_p = None
                parent_container = None
            elif node.name == 'pre':
                is_block = True
                current_p = doc.add_paragraph()
                current_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                parent_container = current_p._element
                raw_text = node.decode_contents()
                add_run(parent_container, raw_text, styles, preserve_newlines=True)
                current_p = None
                parent_container = None
                continue
            elif node.name == 'hr':                                      # Restored: actual page break
                current_p = doc.add_paragraph()
                run = current_p.add_run()
                run._element.append(OxmlElement('w:br', {qn('w:type'): 'page'}))
                current_p = None
                parent_container = None
                continue
            elif node.name == 'b':
                new_styles['bold'] = True
            elif node.name == 'i':
                new_styles['italic'] = True
            elif node.name == 'u':
                new_styles['underline'] = True
            elif node.name == 'a' and 'href' in node.attrs:
                url = node.attrs['href']
                part = doc.part
                r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                if parent_container is None:
                    if current_p is None:
                        current_p = doc.add_paragraph()
                        if justify:
                            current_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        parent_container = current_p._element
                parent_container.append(hyperlink)
                new_styles['hyperlink'] = True
                process_children(node.children, doc, hyperlink, new_styles, preserve_newlines, justify)
                continue
            if not is_block:
                process_children(node.children, doc, parent_container, new_styles, preserve_newlines, justify)

def build_document(input_file, output_file, font_name, font_size, convert_to_pdf, preserve_newlines, justify):
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    soup = BeautifulSoup('<body>' + content + '</body>', 'html.parser')
    body = soup.body
    doc = Document()
    
    # Normal style (body text)
    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)
    
    # Scaled heading sizes (base + hierarchy)
    heading_sizes = {
        1: font_size + 4,
        2: font_size + 2,
        3: font_size,
        4: font_size - 1,
        5: font_size - 2,
        6: font_size - 3,
    }
    for level in range(1, 7):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = font_name
        heading_style.font.size = Pt(heading_sizes[level])
    
    process_children(body.children, doc, preserve_newlines=preserve_newlines, justify=justify)
    
    doc.save(output_file)
    print(f"Word document saved to {output_file}")
    
    if convert_to_pdf:
        pdf_file = output_file.replace('.docx', '.pdf')
        try:
            subprocess.run(['pandoc', output_file, '-o', pdf_file], check=True)
            print(f"PDF saved to {pdf_file}")
        except FileNotFoundError:
            print("Error: pandoc not found. Install it via 'brew install pandoc' on macOS.")
        except subprocess.CalledProcessError as e:
            print(f"Error during PDF conversion: {e}")

def main():
    parser = argparse.ArgumentParser(
        description="Retrnote – Minimalist retro-inspired tag-based word processor (any editor → .docx/PDF)"
    )
    parser.add_argument("input_file", help="Input text file with HTML tags")
    parser.add_argument("output_file", help="Output .docx file")
    parser.add_argument("--font", default="IBM Plex Mono", help="Font name (default: IBM Plex Mono – perfect retro monospace)")
    parser.add_argument("--size", type=int, default=12, help="Font size in points (default: 12)")
    parser.add_argument("--pdf", action="store_true", help="Convert to PDF using pandoc")
    parser.add_argument("--no-preserve-newlines", action="store_true", help="Collapse whitespace (default: preserve newlines)")
    parser.add_argument("--no-justify", action="store_true", help="Left-align paragraphs (default: justify body)")
    
    args = parser.parse_args()
    preserve_newlines = not args.no_preserve_newlines
    justify = not args.no_justify
    build_document(args.input_file, args.output_file, args.font, args.size, args.pdf, preserve_newlines, justify)

if __name__ == "__main__":
    main()
